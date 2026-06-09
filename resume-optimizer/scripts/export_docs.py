"""导出 Word + Markdown 备考包"""
import json, os, sys
from datetime import datetime

INPUT = "answered_questions.json"


def export_markdown(questions):
    lines = [
        "# 面试备考包",
        f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"共 {len(questions)} 题\n",
    ]
    for i, q in enumerate(questions, 1):
        lines.append(f"## {i}. [{q.get('source','')}] {q['question']}")
        if q.get("intent"):
            lines.append(f"\n*面试官在考什么: {q['intent']}*")
        lines.append(f"\n**参考答案:**\n{q.get('answer','')}")
        if q.get("scoring"):
            lines.append("\n**打分标准:**")
            for s in q["scoring"]:
                lines.append(f"- {s.get('dim','')}: 5分={s.get('s5','')} | 3分={s.get('s3','')} | 1分={s.get('s1','')}")
        if q.get("pitfalls"):
            lines.append(f"\n**注意避免:** {'; '.join(q['pitfalls'])}")
        lines.append("\n---\n")

    with open("面试备考包.md", "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"✅ Markdown → 面试备考包.md")


def export_word(questions):
    from docx import Document
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()
    style = doc.styles['Normal']
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    def cn(run):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_heading('面试备考包', level=0)
    doc.add_paragraph(f'{len(questions)} 题 | {datetime.now().strftime("%Y-%m-%d")}')

    for i, q in enumerate(questions, 1):
        doc.add_heading(f'第{i}题 [{q.get("source","")}]', level=2)

        p = doc.add_paragraph()
        r = p.add_run(q['question'])
        cn(r); r.bold = True

        if q.get("intent"):
            p2 = doc.add_paragraph()
            r2 = p2.add_run(f'[考什么] {q["intent"]}')
            cn(r2); r2.italic = True

        p3 = doc.add_paragraph()
        r3 = p3.add_run('【参考答案】')
        cn(r3); r3.bold = True
        p4 = doc.add_paragraph()
        r4 = p4.add_run(q.get('answer', ''))
        cn(r4)

        for s in q.get("scoring", []):
            p5 = doc.add_paragraph()
            r5 = p5.add_run(f'{s.get("dim","")}: 5分-{s.get("s5","")} | 3分-{s.get("s3","")} | 1分-{s.get("s1","")}')
            cn(r5)

        for pf in q.get("pitfalls", []):
            p6 = doc.add_paragraph()
            r6 = p6.add_run(f'  ! {pf}')
            cn(r6); r6.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

        doc.add_paragraph('')

    out = "面试备考包.docx"
    doc.save(out)
    print(f"✅ Word → {out}")


def main():
    if not os.path.exists(INPUT):
        print(f"未找到 {INPUT}，先运行 interview-answer")
        return

    with open(INPUT, encoding="utf-8") as f:
        questions = json.load(f)

    export_markdown(questions)
    export_word(questions)
    print("\n完成！可以分享 面试备考包.docx 或 面试备考包.md")


if __name__ == "__main__":
    main()
